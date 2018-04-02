from docx import Document

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)



path = 'output.docx'








doc = Document(path)

table = doc.tables[0]
for i in reversed(range(1, len(table.rows))):
    row = doc.tables[0].rows[i]
    remove_row(table, row)

for s in student_names:
    for i, entry in enumerate(students[s].keys()):
        row = doc.tables[0].add_row()
        if i == 0:
            row.cells[0].text = s
        
        # f.write('\t' + entry + '\n')
#f.close()


doc.save(path)
